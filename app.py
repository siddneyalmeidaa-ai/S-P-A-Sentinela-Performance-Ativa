import streamlit as st
import pandas as pd
import hashlib
from datetime import datetime
from docx import Document
from docx.shared import Pt
import io

# --- 1. CONFIGURAÇÃO DE AMBIENTE S.A. ---
st.set_page_config(page_title="S.P.A. MASTER - SIDNEY ALMEIDA", layout="wide")

# CSS BLINDADO - LINHA ÚNICA PARA ESTABILIDADE TOTAL NO SERVIDOR
st.markdown("<style>#MainMenu {visibility: hidden;} header {visibility: hidden;} footer {visibility: hidden;} .stDeployButton {display:none;} .manifesto-container {background-color: #050505; border-left: 5px solid #00FF41; padding: 20px; border-radius: 10px; margin-bottom: 25px;} .quote-text { color: #00FF41; font-size: 18px; font-weight: bold; font-style: italic; } .signature { color: #D4AF37; font-size: 12px; text-transform: uppercase; letter-spacing: 2px; } .ofensor-red { color: #FF0000; font-weight: bold; border: 2px solid #FF0000; padding: 10px; border-radius: 5px; text-align: center; }</style>", unsafe_allow_html=True)

# --- 2. BANCO DE DADOS INTEGRAL (PADRÃO OURO S.A.) ---
if 'db' not in st.session_state:
    st.session_state.db = {
        "OPERAÇÃO": {
            "ANA (PERFORMANCE)": {
                "CPF": "123.456.789-01", "VALOR_REAL": 46600.0, "PROJ": 93200.0, 
                "STATUS": "85% LIBERADO", "TEMPO_LOGADO": "08:00:00", "PAUSA": 40, 
                "ALO": 450, "CPC": 120, "CPCA": 95, "PROMESSAS_N": 70, "SABOTAGEM_SCORE": 0
            },
            "MARCOS (SABOTAGEM)": {
                "CPF": "456.123.789-55", "VALOR_REAL": 0.0, "PROJ": 0.0, 
                "STATUS": "0% PENDENTE", "TEMPO_LOGADO": "03:15:00", "PAUSA": 125, 
                "ALO": 12, "CPC": 2, "CPCA": 1, "PROMESSAS_N": 0, "SABOTAGEM_SCORE": 85
            }
        },
        "DISCADOR": {
            "IA_SENTINELA": "ATIVO", "MAILING": "MAILING_OURO_V1", 
            "DESCONHECIDOS": 42.5, "INEXISTENTES": 28.1, "CAIXA_POSTAL": 15.4,
            "VACUO_DETECTADO": False
        },
        "TELEFONIA": {
            "LAT": 380, "SERVER": "Vivo Cloud"
        }
    }

# --- 3. LÓGICA DE PERFORMANCE (CONVERSÃO: PROMESSA / CPCA) ---
LIMITE_PAUSA = 45
df_list = []
for k, v in st.session_state.db["OPERAÇÃO"].items():
    prom = v.get("PROMESSAS_N", 0)
    cpca = v.get("CPCA", 0)
    # REGRA: PROMESSA / CPCA
    conv = round((prom / cpca * 100), 1) if cpca > 0 else 0
    df_list.append({
        "OPERADOR": k, "CPF": v.get("CPF"), "ALÔ": v.get("ALO"), "CPCA": cpca, 
        "PROMESSAS": prom, "CONV %": conv, "REAL": v.get("VALOR_REAL", 0.0),
        "X (-50%)": (v.get("PROJ", 0.0) * 0.5), "LOGADO": v.get("TEMPO_LOGADO", "00:00:00"), 
        "PAUSA": v.get("PAUSA", 0), "SABOTAGEM": v.get("SABOTAGEM_SCORE", 0),
        "STATUS": v.get("STATUS")
    })
df_audit = pd.DataFrame(df_list)

# --- 4. MOTOR JURÍDICO (ADVERTÊNCIA WORD) ---
def gerar_advertencia_word(nome, cpf, pausa, dano):
    doc = Document()
    doc.add_heading('ADVERTÊNCIA DISCIPLINAR - S.A. COMPLIANCE', 0)
    doc.add_paragraph(f"DATA: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"OPERADOR: {nome} | CPF: {cpf}")
    doc.add_heading('RELATÓRIO DE INFRAÇÃO', level=1)
    doc.add_paragraph(f"Detectada pausa excedente de {pausa} min. Limite Operacional: {LIMITE_PAUSA}m.")
    doc.add_paragraph(f"Dano Financeiro Calculado: R$ {dano:.2f}")
    doc.add_paragraph("\n\n__________________________\nAssinatura do Colaborador")
    doc.add_paragraph("__________________________\nSidney Almeida - Comando S.A.")
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# --- 5. INTERFACE S.A. (7 ABAS COMPLETAS) ---
st.markdown('<div class="manifesto-container"><div class="quote-text">"S.P.A. MASTER - SIDNEY ALMEIDA"</div><div class="signature">👊🚀 — COMANDANTE S.A.</div></div>', unsafe_allow_html=True)

tabs = st.tabs(["👑 01. COCKPIT", "👥 02. GESTÃO CPF", "☎️ 03. DISCADOR", "📡 04. TELEFONIA", "🐍 05. SABOTAGEM", "⚖️ 06. JURÍDICO", "📂 07. EXPORTAÇÃO"])

with tabs[0]: # COCKPIT
    c = st.columns(4)
    c[0].metric("CPCA Total", int(df_audit["CPCA"].sum()))
    c[1].metric("Promessas Total", int(df_audit["PROMESSAS"].sum()))
    c[2].metric("Conv Média %", f"{df_audit['CONV %'].mean():.1f}%")
    c[3].metric("Total X (-50%)", f"R$ {df_audit['X (-50%)'].sum():,.2f}")
    st.subheader("🏁 Tabela da Favelinha")
    st.dataframe(df_audit, use_container_width=True)

with tabs[1]: # GESTÃO CPF
    st.header("👥 Auditoria Individual")
    op_sel = st.selectbox("Selecione para Espelhar:", df_audit["OPERADOR"].tolist(), key="op_select")
    res_ind = df_audit[df_audit["OPERADOR"] == op_sel].iloc[0]
    st.info(f"Terminal CPF: {res_ind['CPF']}")
    st.metric("Meta X Indiv.", f"R$ {res_ind['X (-50%)']:,.2f}")
    st.radio("Ação:", ["ENTRA", "PULA", "NÃO ENTRA"], horizontal=True, key="acao_radio")

with tabs[2]: # DISCADOR
    st.header("☎️ Diagnóstico Discador (IA-SENTINELA)")
    d = st.session_state.db["DISCADOR"]
    st.error(f"Lixo Detectado: {(d['DESCONHECIDOS'] + d['INEXISTENTES']):.1f}%")
    st.write(f"Mailing Ativo: {d['MAILING']}")
    if d["VACUO_DETECTADO"]:
        st.warning("⚠️ ZONA DE VÁCUO (1.00x) DETECTADA!")

with tabs[3]: # TELEFONIA
    st.header("📡 Telefonia")
    st.warning(f"Latência: {st.session_state.db['TELEFONIA']['LAT']}ms")

with tabs[4]: # SABOTAGEM
    st.header("🐍 Score de Sabotagem")
    st.table(df_audit[["OPERADOR", "SABOTAGEM", "PAUSA"]])

with tabs[5]: # JURÍDICO
    st.header("⚖️ Gestão de Termos")
    ofensores = df_audit[df_audit["PAUSA"] > LIMITE_PAUSA]
    for _, row in ofensores.iterrows():
        dano_calc = row['PAUSA'] * 0.95
        with st.expander(f"⚖️ GERAR PARA: {row['OPERADOR']}"):
            f_word = gerar_advertencia_word(row['OPERADOR'], row['CPF'], row['PAUSA'], dano_calc)
            st.download_button(f"BAIXAR WORD - {row['OPERADOR']}", f_word, f"Termo_{row['CPF']}.docx", key=f"dl_{row['CPF']}")

with tabs[6]: # EXPORTAÇÃO
    st.header("📂 Exportar Base Ouro")
    csv = df_audit.to_csv(index=False).encode('utf-8-sig')
    st.download_button("EXPORTAR RELATÓRIO FINAL", csv, "S_A_V114.csv", key="export_final")
              
